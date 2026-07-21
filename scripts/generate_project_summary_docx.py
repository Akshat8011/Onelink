"""Generate OneLink Project Summary Word document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date

OUT = r"c:\Users\DELL\Onelink\OneLink_Project_Summary.docx"


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("OneLink\nSmart City Super App")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Complete Project Documentation")
    sr.font.size = Pt(16)
    sr.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"\nLucknow, India\nGenerated: {date.today().strftime('%B %d, %Y')}\n")

    doc.add_page_break()

    # Table of contents
    add_heading(doc, "Table of Contents", 1)
    toc = [
        "1. Executive Summary",
        "2. Vision & Problem Statement",
        "3. Target Audiences",
        "4. High-Level Architecture",
        "5. Repository Structure",
        "6. Technology Stack",
        "7. How a Card Tap Flows (Core Integration)",
        "8. Software Deep Dive",
        "9. Hardware Deep Dive",
        "10. Feature-by-Feature Breakdown",
        "11. MongoDB Data Model",
        "12. MQTT Topic Map",
        "13. Deployment Architecture",
        "14. Development Journey",
        "15. End-to-End Example Scenarios",
        "16. Security Model",
        "17. Resilience & Fallbacks",
        "18. What Makes This Project Distinct",
        "19. Current Git History (Major Milestones)",
        "20. Future Enhancements",
        "21. Quick Reference — Live Endpoints",
    ]
    for item in toc:
        doc.add_paragraph(item)
    doc.add_page_break()

    # 1
    add_heading(doc, "1. Executive Summary", 1)
    add_para(doc, "OneLink is a Lucknow-focused smart-city super app that unifies a digital wallet (NFC/RFID card + in-app balance), supermarket shopping (in-app cart → pay at physical kiosk), metro transit (book tickets, QR entry/exit, slab fares), smart parking (5-zone live grid, allocate, vacate, receipts), and city services (events, bills, EV chargers, vehicle info).")
    add_para(doc, "The differentiator is hardware–software integration: a user taps an ESP32 RFID reader on a Raspberry Pi kiosk, and the mobile web app updates in real time without refresh — wallet debited, orders/receipts synced, parking spots freed.")
    add_table(doc, ["Layer", "Technology", "Live URL"], [
        ["Mobile / Web App", "Expo React Native (web export)", "https://onelink-wine-psi.vercel.app"],
        ["Cloud Backend", "Node.js + Express + Socket.IO + MQTT", "https://onelink-fkqd.onrender.com"],
        ["Database", "MongoDB Atlas", "onelink database"],
        ["Pi Kiosk UI", "Static HTML/JS (local + Vercel)", "http://127.0.0.1:8080 (Pi) or /kiosk (Vercel)"],
        ["Edge Hardware", "ESP32 + RFID + MQTT + brain.py", "On-site at kiosk"],
    ])

    # 2
    add_heading(doc, "2. Vision & Problem Statement", 1)
    add_heading(doc, "Problem", 2)
    add_bullets(doc, [
        "Urban mobility and payments are fragmented across separate apps.",
        "Physical terminals (RFID gates, parking barriers) rarely sync with a user's phone in real time.",
    ])
    add_heading(doc, "Solution", 2)
    add_bullets(doc, [
        "One account — one wallet, one RFID card, one app.",
        "One backend — MongoDB as source of truth.",
        "One physical touchpoint — Pi kiosk for card tap payments.",
        "Live sync — Socket.IO pushes receipts, parking updates, and order status instantly.",
    ])
    add_para(doc, "Geographic focus: Lucknow Metro (UPMRC Red Line) stations from CCS Airport to Munshipulia, local EV chargers (OpenChargeMap), and city events.")

    # 3
    add_heading(doc, "3. Target Audiences", 1)
    add_table(doc, ["Audience", "What They Use", "Value"], [
        ["Daily commuters", "Transit tab, RFID card, metro tickets", "Book trips, tap at gates, track journeys"],
        ["Shoppers", "Shop → Pay via Card → Pi kiosk", "Push cart from phone, pay by tap, receipt in Order History"],
        ["Drivers / parkers", "Parking tab, Pi vacate flow", "Live spots A–E, one spot per user, vacate at kiosk"],
        ["Event-goers", "City tab", "Browse Lucknow events, pay bills"],
        ["EV owners", "Parking → EV tab", "Find chargers via OpenChargeMap"],
        ["Vehicle owners", "Vehicle Info screen", "RTO-style plate lookup"],
        ["Kiosk operators", "Pi display + MQTT gateway", "Physical POS for Shop / Transit / Parking"],
        ["Developers / demo users", "Demo mode fallback", "App works offline with mock data"],
    ])

    # 4
    add_heading(doc, "4. High-Level Architecture", 1)
    add_para(doc, "The system spans four deployable units connected through MongoDB, REST APIs, Socket.IO, MQTT, and a local WebSocket bridge on the Raspberry Pi.")
    add_heading(doc, "Design Principle: Split Responsibilities", 2)
    add_table(doc, ["Component", "Role"], [
        ["Render", "Long-lived server: business logic, Socket.IO, MQTT gateway, authenticated APIs"],
        ["Vercel", "Static app hosting + serverless functions that write MongoDB and notify Render for sockets"],
        ["Raspberry Pi", "Local MQTT→WebSocket bridge; serves kiosk UI over HTTP to avoid mixed-content issues"],
        ["MongoDB Atlas", "Single source of truth for users, wallets, transactions, tickets, carts, parking"],
    ])
    add_para(doc, "Data flow summary: ESP32 → MQTT → brain.py → WebSocket → Kiosk UI → Render API → MongoDB → Socket.IO → Mobile App")

    # 5
    add_heading(doc, "5. Repository Structure", 1)
    add_para(doc, "Onelink/ (monorepo)")
    add_bullets(doc, [
        "backend/ — TypeScript Express API deployed on Render (routes, services, models, MQTT gateway)",
        "mobile/ — Expo app + Vercel deploy root (screens, Zustand stores, api/, public/kiosk/)",
        "hardware/pi/ — brain.py, systemd services, setup-kiosk.sh",
        "scraper/ — BookMyShow event scraper (GitHub Actions daily CI)",
        "quickcomm/ — Separate grocery aggregator sub-project (not core OneLink runtime)",
    ])

    # 6
    add_heading(doc, "6. Technology Stack", 1)
    add_heading(doc, "Backend (Render)", 2)
    add_bullets(doc, [
        "Node.js 18+, Express 4, TypeScript compiled to dist/",
        "Mongoose 7 → MongoDB Atlas",
        "JWT + bcrypt authentication",
        "Socket.IO 4 (user rooms by userId)",
        "MQTT 5 for IoT gateway",
        "Winston logging, Helmet security",
    ])
    add_heading(doc, "Mobile App (Vercel)", 2)
    add_bullets(doc, [
        "Expo ~50, React Native 0.73, react-native-web",
        "React Navigation (tabs + stack)",
        "Zustand state management",
        "Axios → Render /api/v1/*",
        "socket.io-client for realtime",
        "AsyncStorage for JWT token",
        "Build: expo export -p web → dist/",
    ])
    add_heading(doc, "Hardware", 2)
    add_bullets(doc, [
        "ESP32 + RFID reader → MQTT publisher",
        "Raspberry Pi with labwc desktop, 3.5\" kiosk display",
        "Python 3: paho-mqtt, websockets, requests",
        "Chromium kiosk → http://127.0.0.1:8080",
    ])
    add_heading(doc, "Vercel Serverless", 2)
    add_bullets(doc, [
        "hardware-tap.js, pair-card.js, kiosk/check-card.js, kiosk/pay.js",
        "Direct MongoDB access + notify-backend.js → Render internal/notify",
    ])

    # 7
    add_heading(doc, "7. How a Card Tap Flows (Core Integration)", 1)
    steps = [
        "User holds RFID card on ESP32 reader.",
        "ESP32 publishes JSON to MQTT topic onelink/hardware/tap with uid and node.",
        "brain.py on Pi receives MQTT and broadcasts card_tap via WebSocket (ws://127.0.0.1:8765).",
        "Kiosk app.js calls POST /api/v1/kiosk/check-card with cardUid.",
        "If not paired: pairing UI with 10-digit code → POST /api/pair-card → MongoDB → Socket.IO card:paired.",
        "If paired: home screen with Shop, Transit, Parking options.",
        "User selects service, confirms, taps card again.",
        "Kiosk calls Render endpoint (shop/pay, transit/book, parking/allocate or exit).",
        "Backend debits wallet, writes MongoDB, emits Socket.IO events.",
        "Mobile app receives events and refreshes wallet, orders, parking, receipts.",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}", style="List Number")
    add_para(doc, "Why Pi serves kiosk locally: HTTPS (Vercel) blocks ws://127.0.0.1 (mixed content). Local HTTP on port 8080 avoids this.")

    # 8
    add_heading(doc, "8. Software Deep Dive", 1)
    add_heading(doc, "8.1 Backend API Routes", 2)
    add_table(doc, ["Mount", "Purpose"], [
        ["/api/v1/auth", "Register, login, profile, pairing tokens, NFC settings"],
        ["/api/v1/wallet", "Dashboard, top-up, transactions, redeem points"],
        ["/api/v1/transit", "Book ticket, list tickets, journey history"],
        ["/api/v1/mobility", "Parking status, reserve, parking/exit, EV data"],
        ["/api/v1/retail", "Products (FakeStoreAPI), in-app checkout"],
        ["/api/v1/city", "Events, bills, bill payment"],
        ["/api/v1/kiosk", "Pi POS APIs — check-card, shop, transit, parking"],
        ["/api/v1/internal", "POST /notify — Vercel → Socket.IO bridge"],
    ])
    add_heading(doc, "8.2 Kiosk Service", 2)
    add_bullets(doc, [
        "Shop: pushShopCart → KioskCart PENDING → payShopCart → PAID → shop:order-paid",
        "Transit: bookTransitTicket with slab fare → MetroTicket + QR → useTransitTicket for gates",
        "Parking: allocateParking → processEntry → exitParking → ParkingReceipt",
        "Enriched spots with occupant names, timers, demo variety seeding",
    ])
    add_heading(doc, "8.3 Mobile Stores (Zustand)", 2)
    add_table(doc, ["Store", "Responsibility"], [
        ["useAuthStore", "Login, token, profile, master Socket.IO listener"],
        ["useWalletStore", "Balance, transactions, dashboard"],
        ["useRetailStore", "Products, cart, pushCartToKiosk()"],
        ["useOrdersStore", "Order history — syncs kiosk paid carts from API"],
        ["useTransitStore", "Stations, book ticket, active journey"],
        ["useMobilityStore", "Parking spots, receipts, one spot per user"],
        ["useTicketsStore", "Metro ticket receipts with QR"],
    ])
    add_heading(doc, "8.4 Realtime Socket.IO Events", 2)
    add_bullets(doc, [
        "card:paired, shop:cart-pushed, shop:order-paid",
        "transit:entry, transit:exit, transit:ticket-booked",
        "parking:update, parking:entry, parking:exit, parking:receipt",
        "payment:receipt",
    ])

    # 9
    add_heading(doc, "9. Hardware Deep Dive", 1)
    add_heading(doc, "9.1 ESP32 + RFID Reader", 2)
    add_bullets(doc, [
        "Reads card UID (e.g. 72706D05)",
        "Publishes to MQTT topic onelink/hardware/tap",
        "Payload: { uid, node } where node is terminal id",
    ])
    add_heading(doc, "9.2 brain.py (Raspberry Pi)", 2)
    add_bullets(doc, [
        "MQTT subscriber → WebSocket broadcaster on port 8765",
        "Broadcasts { event: card_tap, cardUid, terminalId }",
        "Optional AUTO_PAY=1 legacy direct POST to Vercel hardware-tap",
        "Env: MQTT_BROKER, MQTT_TOPIC, KIOSK_WS_HOST/PORT, HARDWARE_TAP_API_KEY",
    ])
    add_heading(doc, "9.3 Pi Kiosk Setup", 2)
    add_bullets(doc, [
        "Python venv at ~/onelink-venv",
        "brain.py + kiosk files in ~/kiosk/",
        "onelink-brain.service + onelink-kiosk-ui.service (HTTP :8080)",
        "labwc autostart: chromium --kiosk http://127.0.0.1:8080/",
    ])
    add_heading(doc, "9.4 Kiosk UI State Machine", 2)
    add_para(doc, "Steps: idle → checking → pairing | home → service menus → processing → result")
    add_para(doc, "Parking: grid zones A–E, tap your occupied spot → parking_vacate → tap card → POST /kiosk/parking/exit")

    # 10
    add_heading(doc, "10. Feature-by-Feature Breakdown", 1)
    add_heading(doc, "10.1 Authentication & Card Pairing", 2)
    add_bullets(doc, [
        "Register with username + password → JWT",
        "Profile reveals 10-digit pairingToken (password-protected)",
        "Unpaired card at kiosk → pairing keypad → POST /api/pair-card",
        "Socket.IO card:paired → app updates instantly",
    ])
    add_heading(doc, "10.2 Shop (Pay via Card)", 2)
    add_para(doc, "App: Add items → Pay via Card → POST /kiosk/shop/push-cart → KioskCart PENDING")
    add_para(doc, "Pi: Tap card → Shop → select cart → tap → POST /kiosk/shop/pay → wallet debited")
    add_para(doc, "App: GET /kiosk/shop/orders → receipt in Order History")
    add_heading(doc, "10.3 Transit (Metro)", 2)
    add_para(doc, "Stations: CCS Airport, Amausi, Krishna Nagar, Transport Nagar, Alambagh, Charbagh, Hazratganj, Sachivalaya, IT College, Munshipulia")
    add_table(doc, ["Station gap", "Fare (₹)"], [
        ["1", "10"], ["2", "15"], ["3–6", "20"], ["7–9", "30"],
        ["10–13", "40"], ["14–17", "50"], ["18+", "60"],
    ])
    add_heading(doc, "10.4 Parking", 2)
    add_bullets(doc, [
        "20 spots: zones A–E × 4 spots each",
        "Statuses: FREE (green), RESERVED (yellow), OCCUPIED (red)",
        "One spot per user — cannot reserve second until vacating at Pi",
        "App: view-only active session; vacate only at kiosk",
        "Pricing: ratePerMinute (default ₹50) × duration (min 1 minute)",
        "Receipts in Parking → Receipts & Analysis tab",
    ])
    add_heading(doc, "10.5 Wallet", 2)
    add_bullets(doc, [
        "Dashboard: balance, card, transactions, loyalty points",
        "Top-up, redeem points (10 pts = ₹1)",
        "Every payment creates Transaction with category",
    ])
    add_heading(doc, "10.6 City, EV, Vehicle", 2)
    add_bullets(doc, [
        "Events: scraped BookMyShow + MongoDB Event model",
        "Bills: pay utilities from City tab",
        "EV: OpenChargeMap API in Parking → EV tab",
        "Vehicle: plate lookup screen (RTO-style)",
    ])

    # 11
    add_heading(doc, "11. MongoDB Data Model", 1)
    add_table(doc, ["Collection", "Key Purpose"], [
        ["users", "Account, wallet, card, pairing token, activeParkingSpot"],
        ["transactions", "All debits/credits with balance before/after"],
        ["kioskcarts", "Shop carts pushed from app (PENDING → PAID)"],
        ["metrotickets", "Booked tickets with QR payload and status"],
        ["metrojourneys", "Gate tap journeys (MQTT path)"],
        ["parkingspots", "Live grid state per spot"],
        ["parkingreceipts", "Completed parking sessions"],
        ["events", "City events"],
        ["bills", "Utility bills"],
        ["pending_links", "Unpaired card taps (Vercel only)"],
    ])

    # 12
    add_heading(doc, "12. MQTT Topic Map (Backend Gateway)", 1)
    add_heading(doc, "Subscribed (hardware → cloud)", 2)
    add_bullets(doc, [
        "onelink/transit/tap — metro gate taps",
        "onelink/parking/entry|exit|status — parking hardware",
        "onelink/payment/result — payment terminal feedback",
        "onelink/device/heartbeat — device health",
        "onelink/card/pair — pairing at hardware",
    ])
    add_heading(doc, "Published (cloud → hardware)", 2)
    add_bullets(doc, [
        "onelink/transit/gate/command — OPEN/DENY + LED/buzzer",
        "onelink/parking/reserve — LED yellow",
        "onelink/parking/barrier/command — barrier control",
        "onelink/payment/request — NFC payment request",
    ])

    # 13
    add_heading(doc, "13. Deployment Architecture", 1)
    add_table(doc, ["Service", "Platform", "Build", "URL"], [
        ["Backend", "Render", "npm run build (tsc)", "onelink-fkqd.onrender.com"],
        ["Mobile + APIs", "Vercel", "expo export + copy kiosk", "onelink-wine-psi.vercel.app"],
        ["Database", "MongoDB Atlas", "—", "Cloud cluster"],
        ["Pi services", "systemd on Raspberry Pi", "manual setup", "localhost only"],
    ])
    add_heading(doc, "Required Environment Variables", 2)
    add_table(doc, ["Where", "Variables"], [
        ["Render", "MONGODB_URI, JWT_SECRET, INTERNAL_NOTIFY_KEY, MQTT_*, SOCKET_CORS_ORIGIN"],
        ["Vercel", "MONGODB_URI, MONGODB_DB_NAME, HARDWARE_TAP_API_KEY, INTERNAL_NOTIFY_KEY, BACKEND_NOTIFY_URL"],
        ["Pi", "MQTT broker, optional HARDWARE_TAP_API_KEY"],
    ])

    # 14
    add_heading(doc, "14. Development Journey", 1)
    add_heading(doc, "Phase 0 — Foundation", 2)
    add_para(doc, "Expo mobile app, Express backend on Render, MongoDB, JWT auth, demo mode fallbacks.")
    add_heading(doc, "Phase 1 — Hardware Integration", 2)
    add_bullets(doc, [
        "Fixed 405 on Pi POST (Vercel SPA rewrite → serverless hardware-tap API)",
        "brain.py WebSocket bridge, Pi local HTTP :8080",
        "Pairing UI + pair-card API",
    ])
    add_heading(doc, "Phase 2 — Full POS", 2)
    add_bullets(doc, [
        "Kiosk UI → Render /api/v1/kiosk/*",
        "kiosk.routes.ts + kiosk.service.ts",
        "Socket.IO via internal/notify",
        "Mobile Pay via Card, transit sync, parking realtime",
    ])
    add_heading(doc, "Phase 2+ — Parking & Polish", 2)
    add_bullets(doc, [
        "5-zone grid, occupant names, live timers, receipts analysis",
        "Kiosk vacate flow (tap your spot → pay → leave)",
        "One spot per user; removed app vacate button",
        "Shop order history sync from kiosk payments",
        "Kiosk blank page fix, Babel build fix",
    ])

    # 15
    add_heading(doc, "15. End-to-End Example Scenarios", 1)
    add_heading(doc, "Scenario A: Shop via Kiosk", 2)
    add_bullets(doc, [
        "User adds groceries in app, taps Pay via Card",
        "At Pi: tap card → Shop → select cart → tap again",
        "Wallet debited; order in Order History; balance updates live",
    ])
    add_heading(doc, "Scenario B: Park and Vacate", 2)
    add_bullets(doc, [
        "Pi: Parking → select A1 → tap card → spot turns red",
        "App shows active session; cannot book second spot",
        "Later: tap card → Parking → tap A1 → vacate → tap card → charged → receipt in app",
    ])
    add_heading(doc, "Scenario C: Metro Ticket", 2)
    add_bullets(doc, [
        "App books Hazratganj → Munshipulia with QR ticket",
        "Kiosk: Transit → select ticket → tap for entry",
        "Exit gate: tap again → journey completed",
    ])

    # 16
    add_heading(doc, "16. Security Model", 1)
    add_table(doc, ["Layer", "Mechanism"], [
        ["App API", "JWT Bearer on /api/v1/* (except public kiosk endpoints)"],
        ["Kiosk check-card", "Public (card UID only)"],
        ["Vercel hardware-tap", "Bearer HARDWARE_TAP_API_KEY"],
        ["Internal notify", "Bearer INTERNAL_NOTIFY_KEY"],
        ["Pairing", "10-digit token, expiry, one-time use"],
        ["Wallet", "Blocked card check, insufficient balance rejection"],
        ["Parking exit", "Spot ownership validation (occupiedBy === userId)"],
    ])

    # 17
    add_heading(doc, "17. Resilience & Fallbacks", 1)
    add_table(doc, ["Situation", "Behavior"], [
        ["Render down", "Mobile enters demo mode with mock data"],
        ["Socket disconnected", "App works; manual refresh on pull-to-refresh"],
        ["Vercel can't hold WebSockets", "Serverless writes DB → notifies Render for push"],
        ["Pi offline", "App-only flows work; kiosk unavailable"],
        ["MongoDB events empty", "Backend falls back to scraped JSON files"],
    ])

    # 18
    add_heading(doc, "18. What Makes This Project Distinct", 1)
    add_bullets(doc, [
        "True IoT + super-app fusion — not just an app with mock hardware",
        "Dual payment rails — HTTP kiosk APIs + MQTT gate path",
        "Split cloud architecture — Render (stateful) + Vercel (static + burst functions)",
        "Local-first kiosk — Pi serves UI to solve browser security constraints",
        "Realtime by design — every physical action has a digital receipt path",
        "Lucknow-specific — metro stations, EV map, city events grounded in real geography",
    ])

    # 19
    add_heading(doc, "19. Current Git History (Major Milestones)", 1)
    add_table(doc, ["Commit Theme", "What It Added"], [
        ["Hardware-tap Vercel API", "Fix 405 on Pi POST"],
        ["Pi kiosk POS Phase 1", "WebSocket bridge, pairing, basic kiosk"],
        ["Phase 2", "Shop/Transit/Parking kiosk + socket sync"],
        ["Kiosk blank page fix", "Asset paths, parking receipts UI"],
        ["Babel parse fix", "Build stability"],
        ["Kiosk shop orders", "Order history for Pay via Card"],
        ["Parking vacate", "Tap your spot to leave at Pi"],
        ["One spot rule", "No second spot until vacate; no app release button"],
    ])

    # 20
    add_heading(doc, "20. Future Enhancements", 1)
    add_bullets(doc, [
        "ESP32 firmware in repository",
        "Production vehicle RTO API credentials",
        "Official Chalo bus API integration",
        "BookMyShow official API partnership",
        "MQTT on Pi pointing to cloud broker for multi-site",
        "Admin dashboard for device registry / spot management",
        "Push notifications (currently in-app notifications only)",
    ])

    # 21
    add_heading(doc, "21. Quick Reference — Live Endpoints", 1)
    add_table(doc, ["Endpoint", "URL"], [
        ["Render Backend", "https://onelink-fkqd.onrender.com"],
        ["Mobile App", "https://onelink-wine-psi.vercel.app"],
        ["Kiosk (Vercel)", "https://onelink-wine-psi.vercel.app/kiosk"],
        ["Kiosk (Pi)", "http://127.0.0.1:8080"],
        ["WebSocket (Pi)", "ws://127.0.0.1:8765"],
        ["Health Check", "https://onelink-fkqd.onrender.com/health"],
    ])

    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("— End of Document —\nOneLink Smart City Super App")
    fr.italic = True
    fr.font.size = Pt(10)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
