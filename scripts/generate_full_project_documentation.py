"""
Generate OneLink Complete Project Documentation (Word).
Covers architecture, technology choices, hardware, software, bugs, hardening, future.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

OUT = r"c:\Users\DELL\Onelink\docs\OneLink-Complete-Project-Documentation.docx"


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()
    return t


def code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.2)
    return p


def qa(doc, question, answer):
    p = doc.add_paragraph()
    q = p.add_run(f"Q: {question}\n")
    q.bold = True
    p.add_run(f"A: {answer}")
    doc.add_paragraph()


def build():
    doc = Document()
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════════════════════════
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("OneLink\nSmart City Super App")
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Complete Project Documentation\nFrom Concept to Production Hardening")
    sr.font.size = Pt(15)
    sr.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"\nLucknow, India\n"
        f"Document version 1.0\n"
        f"Generated: {date.today().strftime('%B %d, %Y')}\n"
        f"Repository: github.com/Akshat8011/Onelink\n"
        f"Live backend: onelink-fkqd.onrender.com\n"
        f"Live webapp: onelink-wine-psi.vercel.app\n"
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "Table of Contents", 1)
    toc = [
        "Part I — Vision & Problem",
        "  1. Executive Summary",
        "  2. What Problem Does OneLink Solve?",
        "  3. Target Users & Use Cases",
        "Part II — System Architecture",
        "  4. High-Level Architecture",
        "  5. End-to-End Data Flow (Card Tap to Receipt)",
        "  6. Repository Structure",
        "Part III — Technology Choices (The 'Why' Behind Every Decision)",
        "  7. Why Node.js & TypeScript for the Backend?",
        "  8. Why MongoDB Atlas?",
        "  9. Why MQTT for IoT?",
        "  10. Why Socket.IO?",
        "  11. Why Render for the Backend?",
        "  12. Why Vercel for the Web App?",
        "  13. Why Expo / React Native?",
        "  14. Why ESP32 + Raspberry Pi?",
        "  15. Why a Samsung 10.1″ Tablet for the Kiosk?",
        "Part IV — Software Deep Dive",
        "  16. Backend API (Cloud Brain)",
        "  17. Mobile / Web App",
        "  18. Kiosk UI (Point of Sale Terminal)",
        "  19. QuickComm (Price Comparison Module)",
        "  20. Languages, Frameworks & Libraries",
        "Part V — Hardware Deep Dive",
        "  21. Physical Deployment Topology",
        "  22. ESP32 — RFID Reader Role",
        "  23. Raspberry Pi — Edge Bridge Role",
        "  24. Samsung Tablet — Display Role",
        "  25. Hardware ↔ Software Integration",
        "Part VI — Operations, Failures & Safety",
        "  26. Failure Modes & Recovery",
        "  27. Safety Nets & Precautions",
        "  28. Engineering Hardening Sprint (Phases 1–3)",
        "Part VII — Bug History & Lessons Learned",
        "  29. Complete Bug Register",
        "  30. What Each Bug Taught Us",
        "Part VIII — Future & Impact",
        "  31. Verified Proof Results",
        "  32. Future Roadmap",
        "  33. Why OneLink Matters",
        "Appendix A — API Reference Summary",
        "Appendix B — MQTT Topic Map",
        "Appendix C — Environment Variables",
        "Appendix D — Glossary",
    ]
    for line in toc:
        doc.add_paragraph(line)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # PART I
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "Part I — Vision & Problem", 0)

    heading(doc, "1. Executive Summary", 1)
    para(doc,
         "OneLink is a unified smart-city platform built for Lucknow, India. It combines "
         "metro transit, smart parking, retail shopping, bill payments, city events, and "
         "an NFC wallet into a single ecosystem. A user registers on the mobile/web app, "
         "links an RFID smart card, tops up their wallet, and can then pay at physical "
         "kiosks, metro gates, and parking barriers with a single tap.")
    para(doc,
         "The project spans four layers: (1) a cloud backend on Render with MongoDB Atlas, "
         "(2) a cross-platform Expo/React Native app deployed to Vercel, (3) a vanilla-JS "
         "kiosk UI on a Samsung 10.1″ tablet, and (4) ESP32 RFID hardware bridged through "
         "a Raspberry Pi running brain.py. Real-time updates flow via Socket.IO; hardware "
         "events flow via MQTT.")
    bullets(doc, [
        "Production URLs: Backend https://onelink-fkqd.onrender.com · Web https://onelink-wine-psi.vercel.app",
        "14/14 automated behavioural checks pass (transactions, idempotency, rate limits, CORS)",
        "Engineering hardening Phases 1–3 merged to main (July 2026)",
    ])

    heading(doc, "2. What Problem Does OneLink Solve?", 1)
    para(doc,
         "Indian cities suffer from fragmented urban services: separate apps for metro, "
         "parking, shopping, and payments. OneLink unifies these into one wallet and one "
         "card, reducing friction for daily commuters and shoppers.")
    table(doc, ["Pain point", "OneLink solution"],
          [
              ["Multiple payment apps", "Single OneLink wallet + NFC card"],
              ["Cash/QR at metro gates", "Tap card → instant fare debit"],
              ["Unknown parking charges", "Real-time spot tracking + exit charge"],
              ["Shop then pay separately", "Push cart from phone → pay at kiosk"],
              ["No unified transaction history", "All debits in one wallet ledger"],
              ["Hardware silos", "MQTT + Pi bridge connects ESP32 to cloud"],
          ])

    heading(doc, "3. Target Users & Use Cases", 1)
    bullets(doc, [
        "Daily metro commuters — book tickets or tap at gates",
        "Mall/parking users — allocate spot on entry, auto-charge on exit",
        "Shoppers — build cart in mobile app, pay at supermarket kiosk",
        "City explorers — browse BookMyShow events, EV chargers, bus routes",
        "Operators — kiosk at transit hubs, parking lots, retail stores",
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # PART II — ARCHITECTURE
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "Part II — System Architecture", 0)

    heading(doc, "4. High-Level Architecture", 1)
    para(doc, "OneLink follows a hub-and-spoke model: MongoDB Atlas is the source of truth; "
              "the Render backend is the API + MQTT gateway + Socket.IO server; edge devices "
              "publish events upstream and receive commands downstream.")
    code(doc,
         "┌─────────────┐   MQTT    ┌──────────────┐   HTTPS    ┌─────────────────┐\n"
         "│ ESP32 RFID  │ ────────► │ Raspberry Pi │            │ Render Backend  │\n"
         "│  (tap UID)  │           │  brain.py    │            │ Express+Socket  │\n"
         "└─────────────┘           │  WS :8765    │◄──────────►│ MQTT Gateway    │\n"
         "                          └──────┬───────┘            └────────┬────────┘\n"
         "                                 │ WebSocket                      │\n"
         "                          ┌──────▼───────┐            ┌──────────▼────────┐\n"
         "                          │ Samsung Tab  │            │ MongoDB Atlas   │\n"
         "                          │ Kiosk UI     │            │ (users, txns)   │\n"
         "                          └──────────────┘            └─────────────────┘\n"
         "                                                                 ▲\n"
         "                          ┌──────────────┐   HTTPS+WS            │\n"
         "                          │ Vercel Web   │ ──────────────────────┘\n"
         "                          │ Expo App     │\n"
         "                          └──────────────┘")

    heading(doc, "5. End-to-End Data Flow (Card Tap to Receipt)", 1)
    numbered(doc, [
        "User taps RFID card on ESP32 reader at kiosk.",
        "ESP32 reads UID, publishes JSON to MQTT topic onelink/hardware/tap (local broker on Pi).",
        "brain.py on Raspberry Pi receives MQTT message.",
        "brain.py broadcasts {event:'card_tap', cardUid:'A1B2...'} over WebSocket port 8765.",
        "Kiosk browser (app.js) receives WebSocket event, calls POST /api/v1/kiosk/check-card.",
        "Backend looks up user by cardUid (HMAC hash or legacy plaintext), returns balance + pending carts/tickets.",
        "If unregistered: kiosk shows 10-digit pairing keypad → POST /api/pair-card on Vercel.",
        "User selects service (Shop/Transit/Parking), confirms, taps card again.",
        "Kiosk sends payment with Idempotency-Key header to prevent double-charge.",
        "Backend runs MongoDB transaction: debit wallet + write domain record (ticket/cart/receipt) atomically.",
        "Socket.IO pushes payment:receipt to user's mobile app in real time.",
        "Kiosk shows success screen with new balance or insufficient-balance recharge prompt.",
    ])

    heading(doc, "6. Repository Structure", 1)
    table(doc, ["Folder", "Purpose", "Deployed?"],
          [
              ["backend/", "Primary cloud API (Express+TS+MQTT+Socket.IO)", "Yes — Render"],
              ["mobile/", "Expo/React Native app + public/kiosk/ UI", "Yes — Vercel"],
              ["hardware/pi/", "brain.py, systemd services, setup scripts", "On Pi"],
              ["iot-gateway/", "Full serial↔MQTT bridge (metro+parking+ESP32)", "Optional Pi"],
              ["quickcomm/", "Grocery price scraper + Next.js comparison UI", "Separate"],
              ["scraper/", "BookMyShow events scraper (GitHub Actions)", "CI only"],
              ["docs/", "Engineering reports + proof artifacts", "Repo only"],
              ["backend/microservices/", "Legacy scaffold (unused)", "No"],
          ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # PART III — TECHNOLOGY CHOICES
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "Part III — Technology Choices", 0)
    para(doc, "This section answers 'why X?' for every major technology decision.", italic=True)

    heading(doc, "7. Why Node.js & TypeScript for the Backend?", 1)
    qa(doc, "Why not Python/Java/Go?",
       "Node.js shares JavaScript with the kiosk (vanilla JS) and mobile (React Native), "
       "reducing context-switching. Express is battle-tested for REST APIs. TypeScript adds "
       "compile-time safety for payment logic — critical after the ₹13,300 incident.")
    bullets(doc, [
        "Single language across frontend + backend",
        "Excellent async I/O for concurrent MQTT + HTTP + WebSocket",
        "Mongoose ODM integrates natively with MongoDB",
        "Large npm ecosystem (mqtt, socket.io, jsonwebtoken, bcryptjs)",
        "TypeScript catches type errors before deploy (tsc build gate)",
    ])

    heading(doc, "8. Why MongoDB Atlas?", 1)
    qa(doc, "Why not PostgreSQL/MySQL?",
       "User documents are deeply nested (wallet, card, linkedBanks, loyalty). MongoDB's "
       "document model maps naturally. Atlas provides a managed replica set — required for "
       "multi-document ACID transactions introduced in Phase 2.")
    bullets(doc, [
        "Flexible schema for evolving user/card/wallet fields",
        "Replica set on Atlas enables real transactions (debit + txn row atomic)",
        "Free tier sufficient for development; scales horizontally",
        "Mongoose validation + indexes (sparse pairingToken, cardUidHash)",
        "Serverless Vercel functions also read same Atlas cluster",
    ])

    heading(doc, "9. Why MQTT for IoT?", 1)
    qa(doc, "Why not HTTP polling or WebSockets for hardware?",
       "MQTT is the industry standard for IoT pub/sub. ESP32 microcontrollers have tiny "
       "RAM — MQTT clients are lightweight (kilobytes). Pub/sub decouples hardware from "
       "cloud: ESP32 publishes a tap event; brain.py and Render backend both subscribe "
       "independently. QoS 1 guarantees delivery even on flaky Wi-Fi.")
    bullets(doc, [
        "Lightweight protocol — ideal for ESP32 constrained memory",
        "Pub/sub model: one tap → many subscribers (Pi bridge + cloud gateway)",
        "QoS levels for at-least-once delivery",
        "Topic hierarchy: onelink/transit/tap, onelink/parking/entry, onelink/hardware/tap",
        "Brokers: local Mosquitto on Pi (edge), HiveMQ Cloud / EMQX (cloud)",
        "Bidirectional: cloud can publish gate OPEN, barrier UP, LED commands back to hardware",
    ])

    heading(doc, "10. Why Socket.IO?", 1)
    qa(doc, "Why not raw WebSockets everywhere?",
       "Socket.IO adds automatic reconnection, room-based targeting, and fallback to "
       "long-polling. The mobile app joins a room by userId — when a payment completes "
       "at the kiosk, only that user's phone gets the receipt push. Raw WS would require "
       "building all of this from scratch.")
    bullets(doc, [
        "Room-based push: io.to(userId).emit('payment:receipt', ...)",
        "Auto-reconnect with exponential backoff (mobile socket.ts)",
        "Fallback to HTTP polling if WebSocket blocked",
        "Events: transit:entry/exit, parking:update, shop:order-paid, card:paired",
        "Admin room for future dashboard monitoring",
    ])

    heading(doc, "11. Why Render for the Backend?", 1)
    qa(doc, "Why not AWS/GCP/Heroku?",
       "Render offers zero-config Node.js deploys from GitHub with free tier. render.yaml "
       "defines build (npm run build) and start (node dist/server.js) declaratively. "
       "Trade-off: free tier sleeps after inactivity (cold starts) — mitigated by kiosk "
       "keep-alive ping every 4 minutes.")
    bullets(doc, [
        "Git-push deploy from main branch",
        "Built-in HTTPS + custom domains",
        "Environment variable management for secrets",
        "Limitation: single instance (rate limiter is per-process)",
        "Cold start ~15–30s on free tier — kiosk pings /health to keep warm",
    ])

    heading(doc, "12. Why Vercel for the Web App?", 1)
    qa(doc, "Why not host everything on Render?",
       "Vercel excels at static/SPA hosting with global CDN edge. The Expo web export "
       "is a static bundle — Vercel serves it instantly worldwide. Vercel also runs "
       "serverless functions (mobile/api/) for hardware-tap and pair-card endpoints "
       "that need low-latency edge execution close to the Pi.")
    bullets(doc, [
        "Global CDN for mobile web app (onelink-wine-psi.vercel.app)",
        "Serverless API routes: /api/pair-card, /api/hardware-tap, /api/kiosk/*",
        "SPA rewrites in vercel.json — /kiosk serves kiosk index.html",
        "Automatic HTTPS, preview deploys per branch",
        "Pairing endpoint on Vercel because Pi brain.py AUTO_PAY mode POSTs here",
    ])

    heading(doc, "13. Why Expo / React Native?", 1)
    qa(doc, "Why not Flutter or native Swift/Kotlin?",
       "One codebase → iOS, Android, and Web. Expo SDK 50 provides navigation, fonts, "
       "and build tooling. Zustand for lightweight state. The same app runs in browser "
       "(Vercel) and can be built as APK/IPA later.")
    bullets(doc, [
        "Write once, deploy to mobile + web",
        "React Navigation for tab + stack screens",
        "Zustand stores (auth, wallet, mobility, shop)",
        "AsyncStorage for JWT persistence",
        "Socket.IO client for real-time receipts",
        "patch-package for dependency fixes",
    ])

    heading(doc, "14. Why ESP32 + Raspberry Pi?", 1)
    qa(doc, "Why not connect ESP32 directly to the cloud?",
       "ESP32 has limited TLS/HTTPS capability and no browser. The Pi acts as an edge "
       "gateway: runs local MQTT broker, bridges to WebSocket for the kiosk browser, "
       "and can optionally forward to cloud. Pi also runs systemd services for "
       "auto-restart and serves the kiosk HTML on port 8080.")
    table(doc, ["Device", "Role", "Why this device"],
          [
              ["ESP32 + RFID module", "Read NFC card UID, publish to MQTT",
               "Cheap (~₹500), low power, built-in Wi-Fi, huge community, PN532/MFRC522 support"],
              ["Raspberry Pi 4", "MQTT broker, brain.py bridge, HTTP server, systemd",
               "Full Linux, runs Python, local WebSocket server, auto-restart via systemd"],
              ["Samsung Tab 10.1″", "Kiosk touchscreen display (1280×800)",
               "Large readable UI, capacitive touch, runs Chromium kiosk mode"],
          ])

    heading(doc, "15. Why a Samsung 10.1″ Tablet for the Kiosk?", 1)
    bullets(doc, [
        "1280×800 resolution — kiosk CSS viewport tuned exactly for this",
        "Capacitive multi-touch for pairing keypad and service selection",
        "Runs Chromium in kiosk/fullscreen mode pointed at http://<pi-ip>:8080/kiosk",
        "Connects to Pi WebSocket at ws://<pi-ip>:8765 for card tap events",
        "Can also load kiosk from Vercel CDN as fallback",
        "Reader IP configurable via ?reader=<pi-ip> URL parameter or localStorage",
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # PART IV — SOFTWARE
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "Part IV — Software Deep Dive", 0)

    heading(doc, "16. Backend API (Cloud Brain)", 1)
    para(doc, "Entry: backend/src/server.ts — Express on PORT 10000 (Render) or 5000 (dev).")
    table(doc, ["Route prefix", "Service", "Key operations"],
          [
              ["/api/v1/auth", "Authentication", "register, login, pair-card, delink-card, reveal pairing token"],
              ["/api/v1/wallet", "Wallet", "balance, top-up, transactions, redeem points, card settings"],
              ["/api/v1/transit", "Metro", "journey entry/exit, fare calculation, tickets"],
              ["/api/v1/mobility", "Parking", "spot allocation, exit charges, enriched grid"],
              ["/api/v1/retail", "Shopping", "product catalog, order placement"],
              ["/api/v1/city", "Events", "city events from scraper data"],
              ["/api/v1/kiosk", "Kiosk POS", "check-card, shop pay, transit book, parking allocate/exit"],
              ["/api/v1/internal", "Internal", "Vercel→Socket.IO notify bridge"],
          ])
    para(doc, "Core services:", bold=True)
    bullets(doc, [
        "wallet.service.ts — atomic debitInSession(), processPayment() in Mongo transaction",
        "kiosk.service.ts — orchestrates kiosk flows with idempotency + transactions",
        "mqtt-gateway.ts — subscribes to hardware topics, drives gates/barriers/payments",
        "parking.service.ts — 20 spots (zones A–E), rate per minute, exit payment guard",
        "transit.service.ts — Lucknow metro 10-station matrix, slab + per-km fares",
        "reversal.service.ts — sweeps stale PENDING carts/transactions every 5 min",
    ])

    heading(doc, "17. Mobile / Web App", 1)
    para(doc, "Framework: Expo SDK 50 + React Native 0.73. Deployed as static web to Vercel.")
    para(doc, "Navigation structure:", bold=True)
    bullets(doc, [
        "Auth gate: LoginScreen (username+password or card+CVV)",
        "Bottom tabs: Home, Shop, Transit, Parking, City (Events), Wallet",
        "Stack screens: Notifications, Settings, Rewards, Tickets, OrderHistory, Account section",
    ])
    para(doc, "State management: Zustand stores in mobile/src/store/")
    para(doc, "API: axios → https://onelink-fkqd.onrender.com/api/v1 with JWT Bearer token")
    para(doc, "Real-time: Socket.IO singleton joins userId room, listens for payment:receipt, transit:*, parking:*, card:paired")
    para(doc, "Integrations (INTEGRATION_SUMMARY.md):", bold=True)
    bullets(doc, [
        "UPMRC metro map — 22 stations, interactive MetroMapComponent",
        "OpenChargeMap — real EV charger API (Lucknow 50km radius)",
        "Chalo Bus — live bus tracking (mock-ready for API key)",
        "Vehicle Info — RTO lookup structure (demo data, needs credentials)",
        "BookMyShow events — scraper → events_lucknow.scraped.json in app bundle",
        "QuickComm — grocery price comparison via FastAPI",
    ])

    heading(doc, "18. Kiosk UI (Point of Sale Terminal)", 1)
    para(doc, "Files: mobile/public/kiosk/index.html, app.js (~1500 lines), styles.css (~1000 lines)")
    para(doc, "Technology: Vanilla JavaScript (no framework) — intentional for Pi performance. No build step, no React overhead on constrained tablet.")
    para(doc, "Screen flow:", bold=True)
    code(doc, "idle → checking → (pairing | home) → service screens → confirm → processing → result")
    bullets(doc, [
        "Idle: animated rings, 'Tap Your Card to Begin' with elegant white wave animation",
        "Home: personalized greeting, balance, loyalty tier, stats row, service cards, low-balance warning",
        "Transit: station picker → fare preview (localSlabFare, no network) → tap to book",
        "Shop: pending carts from mobile app → review → tap to pay",
        "Parking: A–E zone grid → allocate or vacate → tap to confirm",
        "Pairing: 10-digit keypad for new card registration",
        "Help: full user guide with FAQs, webapp URL, step-by-step per service",
        "Performance: delegated click handler, prefetch parking/stations, backend keep-alive ping",
        "Idempotency-Key per payment attempt prevents double-charge on double-tap",
    ])

    heading(doc, "19. QuickComm (Price Comparison Module)", 1)
    para(doc, "Separate sub-project in quickcomm/ — grocery price aggregator across Blinkit, Zepto, Swiggy Instamart using Playwright scrapers. FastAPI backend on port 8001, Next.js 16 web UI on port 3001. Mobile app consumes via quickCommApi.ts with demo fallback.")

    heading(doc, "20. Languages, Frameworks & Libraries", 1)
    table(doc, ["Layer", "Languages", "Key libraries"],
          [
              ["Backend", "TypeScript → JavaScript", "Express, Mongoose, mqtt, socket.io, jsonwebtoken, bcryptjs, helmet, cors, winston"],
              ["Mobile app", "TypeScript/JavaScript + JSX", "Expo, React Native, React Navigation, Zustand, axios, socket.io-client"],
              ["Kiosk UI", "HTML, CSS, Vanilla JS", "None (zero dependencies)"],
              ["Pi brain", "Python 3", "paho-mqtt, websockets, requests, systemd.daemon"],
              ["IoT gateway", "Python 3", "paho-mqtt, pyserial"],
              ["QuickComm", "Python + TypeScript", "Playwright, FastAPI, Next.js 16, Tailwind"],
              ["Scraper", "Python", "Selenium, BeautifulSoup"],
              ["DevOps", "YAML, Shell", "GitHub Actions, systemd, render.yaml, vercel.json"],
          ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # PART V — HARDWARE
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "Part V — Hardware Deep Dive", 0)

    heading(doc, "21. Physical Deployment Topology", 1)
    para(doc, "A typical kiosk deployment at a metro station or supermarket:")
    code(doc,
         "┌─────────────────────────────────────────────────────────────┐\n"
         "│  Kiosk enclosure (counter/desk)                             │\n"
         "│  ┌──────────────┐    USB/serial    ┌──────────────┐        │\n"
         "│  │ Samsung Tab  │◄── WebSocket ──►│ Raspberry Pi │        │\n"
         "│  │ 10.1″ display│    (Wi-Fi LAN)  │  (hidden)    │        │\n"
         "│  │ Chromium     │                 │ Mosquitto    │        │\n"
         "│  │ kiosk mode   │                 │ brain.py     │        │\n"
         "│  └──────────────┘                 │ :8080 HTTP   │        │\n"
         "│         ▲                         │ :8765 WS     │        │\n"
         "│         │ tap                     └──────┬───────┘        │\n"
         "│  ┌──────┴───────┐                        │ MQTT           │\n"
         "│  │ ESP32 + RFID │◄───────────────────────┘                │\n"
         "│  │ reader module│   Wi-Fi to Pi broker                      │\n"
         "│  └──────────────┘                                           │\n"
         "└─────────────────────────────────────────────────────────────┘\n"
         "         │ HTTPS (internet)\n"
         "         ▼\n"
         "   Render Backend + MongoDB Atlas + Vercel Web App")

    heading(doc, "22. ESP32 — RFID Reader Role", 1)
    para(doc, "The ESP32 is a low-cost microcontroller with built-in Wi-Fi. Paired with an "
              "RFID/NFC module (PN532 or MFRC522), it reads the unique identifier (UID) "
              "when a user taps their card.")
    bullets(doc, [
        "Reads 4–16 byte hex UID from NFC card (ISO 14443 Type A)",
        "Publishes JSON to MQTT: { cardUid, deviceId, timestamp }",
        "Topic: onelink/hardware/tap (local broker) or onelink/transit/tap (cloud)",
        "Can also communicate via serial at 115200 baud (iot-gateway protocol)",
        "Serial format: USER:Name|UID:CardUID|ACTION:PAYMENT_SUCCESS|AMOUNT:100|TYPE:SHOPPING",
        "Firmware source is NOT in this repo — only the protocol contract is defined",
        "Receives commands back: GATE:OPEN, BARRIER:UP, LED:GREEN/RED, PAY:amount",
    ])

    heading(doc, "23. Raspberry Pi — Edge Bridge Role", 1)
    para(doc, "brain.py (hardware/pi/brain.py) is the production edge software:")
    numbered(doc, [
        "Subscribes to MQTT topic onelink/hardware/tap on local broker (127.0.0.1:1883)",
        "On each tap: broadcasts {event:'card_tap', cardUid} to all WebSocket clients on :8765",
        "Maintains set of connected kiosk browsers; handles connect/disconnect",
        "MQTT reconnect: exponential backoff 1s→30s with on_disconnect logging",
        "WebSocket server: restart loop on failure (never dies permanently)",
        "Optional sd_notify: READY=1 + WATCHDOG=1 every 15s for systemd watchdog",
        "Optional AUTO_PAY=1: POST directly to Vercel /api/hardware-tap (legacy mode)",
    ])
    para(doc, "systemd services (auto-start on boot):", bold=True)
    table(doc, ["Service", "What it runs", "Restart policy"],
          [
              ["onelink-kiosk-ui.service", "python3 -m http.server 8080 (serves kiosk HTML)", "Restart=always, 3s"],
              ["onelink-brain.service", "python3 brain.py (MQTT→WS bridge)", "Type=notify, WatchdogSec=30"],
          ])

    heading(doc, "24. Samsung Tablet — Display Role", 1)
    bullets(doc, [
        "Runs Chromium in kiosk/fullscreen mode",
        "URL: http://127.0.0.1:8080/ (local Pi) or http://<pi-ip>:8080/?reader=<pi-ip>",
        "Viewport: 1280×800, user-scalable=no (locked for POS use)",
        "WebSocket connects to ws://<reader-ip>:8765 for real-time card tap events",
        "Reader status indicator in header: green pulse = online, red = offline",
        "No app install needed — pure browser, easy to replace/reset",
        "localStorage stores reader IP for reconnect across page loads",
    ])

    heading(doc, "25. Hardware ↔ Software Integration", 1)
    table(doc, ["Layer transition", "Protocol", "Data format"],
          [
              ["Card → ESP32", "NFC/RFID (13.56 MHz)", "4–16 byte UID"],
              ["ESP32 → Pi", "MQTT (TCP 1883)", "JSON { cardUid, deviceId }"],
              ["Pi → Tablet", "WebSocket (TCP 8765)", "JSON { event:'card_tap', cardUid }"],
              ["Tablet → Cloud", "HTTPS REST", "JSON + Idempotency-Key header"],
              ["Cloud → Mobile", "Socket.IO (WSS)", "JSON events per userId room"],
              ["Cloud → ESP32", "MQTT publish", "GATE:OPEN, BARRIER:UP, LED commands"],
              ["Pi → Cloud (legacy)", "HTTPS POST", "Vercel /api/hardware-tap"],
          ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # PART VI — FAILURES & SAFETY
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "Part VI — Operations, Failures & Safety", 0)

    heading(doc, "26. Failure Modes & Recovery", 1)
    table(doc, ["Failure", "Symptom", "Recovery mechanism", "Measured?"],
          [
              ["ESP32 Wi-Fi drop", "Taps not detected", "ESP32 firmware reconnect (gap: firmware not in repo)", "Pi test pending"],
              ["MQTT broker down", "brain.py logs disconnect", "Exponential backoff reconnect 1s→30s", "Pi test pending"],
              ["brain.py crash", "Reader Offline on kiosk", "systemd Restart=always within ~3s", "Pi test pending"],
              ["WebSocket disconnect", "Reader Offline", "Kiosk exponential backoff reconnect 1s→30s", "Pi test pending"],
              ["Render cold start", "Kiosk slow first tap", "Keep-alive ping every 4 min to /health", "Mitigated"],
              ["MongoDB disconnect", "API errors", "Mongoose auto-reconnect + server logs", "Atlas managed"],
              ["Double-tap payment", "Risk of double charge", "Idempotency-Key dedup", "PROVEN: 1 debit"],
              ["Crash mid-payment", "Partial DB state", "MongoDB transaction rollback", "PROVEN: 0 orphan rows"],
              ["Insufficient balance", "False receipt", "Block before domain write + UI recharge prompt", "PROVEN: 4935→4935"],
              ["Abandoned cart", "PENDING forever", "Reversal sweeper → CANCELLED after 30 min", "PROVEN"],
              ["Brute-force pairing", "Token guessing", "Rate limit 10/min per IP", "PROVEN: 429 at #11"],
          ])

    heading(doc, "27. Safety Nets & Precautions", 1)
    bullets(doc, [
        "MongoDB multi-document transactions — debit + record commit or rollback together",
        "Idempotency keys — duplicate requests return cached result, never re-charge",
        "Balance re-read inside transaction — prevents TOCTOU race on concurrent taps",
        "paymentResult.success check on every caller — no operation proceeds on failed payment",
        "Card blocked check before every debit",
        "Rate limiting on unauthenticated endpoints (check-card, pair-card, login)",
        "Input validation — type:'string' blocks NoSQL operator injection",
        "CORS allowlist — only known origins (vercel, digitalzen, localhost, LAN)",
        "HMAC card UID hashing — dual-read migration, plaintext eventually dropped",
        "systemd auto-restart for Pi services",
        "Reversal sweeper for stale pending state",
        "JWT authentication for mobile API routes",
        "Helmet security headers on Express",
        "Password required to reveal pairing token",
        "10-digit pairing token with 7-day expiry",
        "One parking spot per user enforced",
    ])

    heading(doc, "28. Engineering Hardening Sprint (Phases 1–3)", 1)
    table(doc, ["Phase", "Focus", "Key deliverables", "Status"],
          [
              ["Phase 1", "Self-healing", "brain systemd, MQTT/WS reconnect, Socket.IO status API", "Merged + Pi proof pending"],
              ["Phase 2", "Transactional integrity", "Mongo transactions, idempotency, reversal sweeper, validation", "Merged + 6/6 proven"],
              ["Phase 3", "Security", "Rate limits, HMAC card hashing, CORS tighten", "Merged + 8/8 proven"],
              ["Phase 4", "Documentation", "FMEA, runbooks, API contract", "Planned"],
          ])
    para(doc, "Full details: docs/OneLink-Engineering-Hardening-Report.md")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # PART VII — BUGS
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "Part VII — Bug History & Lessons Learned", 0)

    heading(doc, "29. Complete Bug Register", 1)
    para(doc, "Chronological register of significant bugs found and fixed during development, "
              "derived from git history and engineering sessions.", italic=True)

    bugs = [
        ("B01", "Card pairing 500 error", "Valid 10-digit code → 'Card pairing failed' HTTP 500",
         "Multiple causes: non-sparse unique index on pairingToken (null collision), duplicate cardUid, full-document validation on legacy users",
         "Sparse+unique index rebuild; unique consumed marker used_<userId> instead of null; dot-notation updates; UID reclaim from stale accounts",
         "Always use sparse indexes for nullable unique fields; never user.save() on legacy partial docs"),
        ("B02", "Double digit on Pi keypad", "Tapping pairing keypad registers digit twice",
         "Touch events firing both touchstart and click",
         "Single delegated click handler; touch-action: manipulation; debounce",
         "Test on actual touchscreen hardware, not just mouse"),
        ("B03", "Double-tap opens pairing twice", "Rapid card tap triggers duplicate pairing flow",
         "No debounce on card_tap WebSocket handler",
         "State guard: ignore tap while step !== idle/awaiting",
         "Hardware events need software debouncing"),
        ("B04", "Kiosk blank page", "Kiosk shows white screen on Pi",
         "JavaScript error in render path; missing DOM element",
         "Defensive null checks; fixed render guard conditions",
         "Always test with node --check and on actual Pi browser"),
        ("B05", "Babel parse error ?? and ||", "Mobile app build fails",
         "Mixed nullish coalescing and logical OR without parentheses",
         "Added explicit parentheses in useMobilityStore",
         "Babel requires parentheses when mixing ?? and ||"),
        ("B06", "Shop payments not in order history", "Kiosk shop pay succeeded but app shows no order",
         "Kiosk used separate flow; order history only read mobile-placed orders",
         "Sync kiosk PAID carts to shop orders API; emit shop:order-paid event",
         "All payment paths must write to same order collection"),
        ("B07", "Pi POST 405 on hardware-tap", "brain.py AUTO_PAY POST returns Method Not Allowed",
         "Render backend had no POST route for hardware tap; needed serverless",
         "Added Vercel serverless /api/hardware-tap route",
         "Edge devices need dedicated serverless endpoints, not monolith routes"),
        ("B08", "Insufficient balance false receipt (₹13,300 on ₹4,935)", "Payment 'succeeded' with receipt despite low balance",
         "parking.service processExit and transit.service processExit ignored paymentResult.success; retail route returned fabricated receipt; MQTT gateway swallowed failed payments",
         "Check success before freeing spot/completing journey; HTTP 402 on retail; emit FAILED receipt on MQTT; kiosk showFailure() with recharge prompt",
         "NEVER proceed after processPayment without checking success — this is the most critical lesson"),
        ("B09", "Kiosk UI lag (1–2+ seconds per tap)", "Screen transitions feel sluggish on tablet",
         "Blocking await before render; full re-render + event rebind; backdrop-filter GPU load; Render cold starts",
         "Render-first with cache; delegated click handler; localSlabFare(); removed backdrop-filter; backend keep-alive ping",
         "Measure on target hardware; prefetch; never block UI on network"),
        ("B10", "Basic/unprofessional kiosk UI", "User reported UI as 'pathetic and basic'",
         "Legacy minimal HTML with no branding, info, or animations",
         "Full redesign: gradients, stats, weather, wave animation, user guide, 1280×800 layout",
         "POS UI is the user's first impression — invest in design"),
        ("B11", "Need Help button non-functional", "Button did nothing on tap",
         "Missing click handler / wrong step routing",
         "Added guide screen with FAQs, steps, webapp URL, service-specific instructions",
         "Every interactive element needs a handler and a test on target hardware"),
        ("B12", "Non-atomic payments", "Theoretical: debit without ticket, or ticket without debit",
         "user.save() and Transaction.create() were separate non-transactional writes",
         "runInTransaction() + debitInSession(); composite flows in single transaction",
         "Money paths MUST use database transactions on replica set"),
        ("B13", "Double-charge on retry/double-tap", "Same payment could debit wallet twice",
         "No idempotency mechanism",
         "IdempotencyKey model + withIdempotency() + client UUID per attempt",
         "Every money-moving endpoint needs idempotency keys"),
        ("B14", "Open CORS (security)", "Any website could make credentialed API calls",
         "SOCKET_CORS_ORIGIN unset defaulted to allow *",
         "Explicit allowlist; allow-all requires explicit * opt-in",
         "Secure by default; opt-in to permissive"),
        ("B15", "Plaintext card UIDs in database", "DB leak exposes raw RFID identifiers",
         "cardUid stored as plaintext only",
         "HMAC-SHA256 cardUidHash + dual-read migration path",
         "PII and hardware identifiers should be hashed at rest"),
        ("B16", "brain.py no reconnect", "MQTT drop = manual restart required",
         "connect() + loop_forever() with no on_disconnect",
         "Exponential backoff, WS restart loop, systemd watchdog",
         "Edge software must self-heal — operators won't SSH in"),
        ("B17", "Event images broken", "BookMyShow posters not loading",
         "Wrong image URL format; missing transforms",
         "ImageKit transforms; landscape URL extraction; robust loading fallbacks",
         "Scraped content needs image URL validation pipeline"),
        ("B18", "Shop catalog miscategorization", "Products in wrong categories",
         "Keyword matching too broad",
         "Improved categorization rules and product image matching",
         "Catalog quality affects user trust"),
    ]

    table(doc, ["ID", "Bug", "Symptom", "Root cause", "Fix", "Lesson"],
          [[b[0], b[1], b[2], b[3], b[4], b[5]] for b in bugs])

    heading(doc, "30. What Each Bug Taught Us — Summary", 1)
    bullets(doc, [
        "Payment integrity is non-negotiable — always check success, use transactions, use idempotency",
        "Database indexes matter — sparse for nullable unique fields (pairingToken lesson)",
        "Test on real hardware — touchscreen, Pi, tablet behave differently from desktop",
        "Edge devices must self-heal — systemd, reconnect, watchdog",
        "UI performance is a feature — users perceive lag as broken",
        "Security defaults matter — CORS, rate limits, hashing should be on by default",
        "Legacy data is a trap — partial documents break full save(); use dot-notation updates",
        "One wallet, many paths — every payment entry point must share the same guards",
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # PART VIII — FUTURE
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "Part VIII — Future & Impact", 0)

    heading(doc, "31. Verified Proof Results (July 2026)", 1)
    para(doc, "14/14 automated checks pass against compiled production code:")
    table(doc, ["#", "Test", "Result"],
          [
              ["1", "₹13,300 on ₹4,935 balance", "BLOCKED — balance unchanged, 0 txn rows"],
              ["2", "Atomic debit", "balance 1000→750, 1 txn row"],
              ["3", "Crash rollback", "balance stayed 2000, 0 orphans"],
              ["4", "Sequential double-tap", "1 debit, replayed same txnId"],
              ["5", "Concurrent double-tap", "1 success, 1 debit"],
              ["6", "Reversal sweep", "stale CANCELLED, fresh PENDING"],
              ["7", "Rate limit check-card", "first 429 at request #31"],
              ["8", "Rate limit per-IP", "independent buckets"],
              ["9", "429 response shape", "Retry-After=60"],
              ["10", "CORS allowlist", "9/9 correct"],
              ["11", "NoSQL injection block", "400 must be string"],
              ["12", "Card UID HMAC", "deterministic 64-hex"],
              ["13", "Dual-read query", "hash OR plaintext"],
              ["14", "Build + syntax", "tsc exit 0, node --check pass"],
          ])
    para(doc, "Reproduce: cd backend && npm run verify")

    heading(doc, "32. Future Roadmap", 1)
    bullets(doc, [
        "Phase 4: FMEA document, incident runbooks, API contract specification",
        "ESP32 firmware in repo with OTA update support",
        "Drop plaintext cardUid after Atlas migration completes",
        "Redis-backed rate limiter for horizontal scaling",
        "Native mobile builds (APK/IPA) via Expo EAS",
        "Real BookMyShow / Chalo / Vehicle Info API partnerships",
        "UPI integration for wallet top-up",
        "Admin dashboard with live device monitoring",
        "Multi-terminal kiosk fleet management",
        "Offline kiosk queue (display-only, no payment) for network outages",
        "PCI-DSS considerations if handling real card networks",
        "Lucknow metro gate integration (physical turnstile MQTT)",
    ])

    heading(doc, "33. Why OneLink Matters", 1)
    para(doc,
         "OneLink demonstrates that a full smart-city payment ecosystem can be built by a "
         "small team using modern cloud-native tools. It connects physical RFID hardware "
         "to a cloud wallet through a Raspberry Pi edge bridge, serves a professional "
         "kiosk UI on a commodity tablet, and pushes real-time receipts to a cross-platform "
         "mobile app — all while maintaining transactional integrity proven by 14 automated checks.")
    bullets(doc, [
        "Unified wallet reduces payment friction for daily commuters",
        "Open architecture (MQTT topics, REST API) allows new services without rewrites",
        "Edge + cloud split keeps kiosk responsive even with cloud latency",
        "Engineering hardening sprint shows path from prototype to production-grade",
        "Replicable model for other Indian cities (metro + parking + retail)",
        "Educational value: full-stack IoT + fintech + mobile in one repo",
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # APPENDICES
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "Appendix A — API Reference Summary", 1)
    table(doc, ["Method", "Endpoint", "Auth", "Purpose"],
          [
              ["POST", "/api/v1/auth/register", "No", "Create account + pairing token"],
              ["POST", "/api/v1/auth/login", "No", "Login (username or card+CVV)"],
              ["POST", "/api/v1/auth/pair-card", "No", "Link RFID UID to account"],
              ["POST", "/api/v1/kiosk/check-card", "No", "Lookup user by card UID"],
              ["POST", "/api/v1/kiosk/shop/pay", "No", "Pay pending cart at kiosk"],
              ["POST", "/api/v1/kiosk/transit/book", "No", "Book metro ticket"],
              ["POST", "/api/v1/kiosk/parking/allocate", "No", "Allocate parking spot"],
              ["POST", "/api/v1/kiosk/parking/exit", "No", "Exit parking + charge"],
              ["GET", "/api/v1/wallet", "JWT", "Wallet dashboard"],
              ["POST", "/api/v1/wallet/top-up", "JWT", "Add money to wallet"],
              ["GET", "/health", "No", "Server health + MQTT device status"],
          ])

    heading(doc, "Appendix B — MQTT Topic Map", 1)
    table(doc, ["Topic", "Direction", "Purpose"],
          [
              ["onelink/hardware/tap", "Device→Cloud", "ESP32 NFC tap (Pi brain)"],
              ["onelink/transit/tap", "Device→Cloud", "Metro gate tap (entry/exit)"],
              ["onelink/transit/gate/command", "Cloud→Device", "OPEN/DENY/PROMPT_PAIR"],
              ["onelink/parking/entry", "Device→Cloud", "Vehicle entering spot"],
              ["onelink/parking/exit", "Device→Cloud", "Vehicle leaving spot"],
              ["onelink/parking/barrier/command", "Cloud→Device", "RAISE/LOWER barrier"],
              ["onelink/payment/result", "Device→Cloud", "NFC payment outcome"],
              ["onelink/payment/request", "Cloud→Device", "Request payment on reader"],
              ["onelink/card/pair", "Device→Cloud", "Pairing request from terminal"],
              ["onelink/card/pair/result", "Cloud→Device", "Pairing success/failure"],
              ["onelink/device/heartbeat", "Device→Cloud", "Device alive signal"],
          ])

    heading(doc, "Appendix C — Environment Variables", 1)
    table(doc, ["Variable", "Where", "Purpose"],
          [
              ["MONGODB_URI", "Render", "Atlas connection string"],
              ["JWT_SECRET", "Render", "Auth token signing"],
              ["CARD_UID_HMAC_SECRET", "Render", "HMAC key for card UID hashing"],
              ["MQTT_BROKER_URL", "Render", "Cloud MQTT broker URL"],
              ["SOCKET_CORS_ORIGIN", "Render", "CORS allowlist (* for open)"],
              ["PORT", "Render", "HTTP port (10000)"],
              ["MQTT_BROKER", "Pi", "Local broker (127.0.0.1)"],
              ["KIOSK_WS_HOST", "Pi", "WebSocket bind (0.0.0.0)"],
              ["AUTO_PAY", "Pi", "Legacy direct payment mode"],
              ["HARDWARE_TAP_API_KEY", "Pi/Vercel", "Bearer key for hardware-tap API"],
          ])

    heading(doc, "Appendix D — Glossary", 1)
    table(doc, ["Term", "Definition"],
          [
              ["NFC/RFID", "Near Field Communication — contactless card technology"],
              ["UID", "Unique Identifier burned into each RFID card/chip"],
              ["MQTT", "Message Queuing Telemetry Transport — lightweight IoT pub/sub protocol"],
              ["Socket.IO", "Real-time bidirectional event library (WebSocket + fallback)"],
              ["Kiosk", "Fixed-point-of-sale terminal with touchscreen + card reader"],
              ["brain.py", "Raspberry Pi script bridging MQTT card taps to WebSocket"],
              ["Idempotency", "Same request twice = same outcome, no double effect"],
              ["ACID transaction", "Database operation that fully commits or fully rolls back"],
              ["Pairing token", "10-digit code linking a physical card to a digital account"],
              ["Slab fare", "Metro fare based on number of stations traveled, not distance"],
              ["Render cold start", "Free-tier server sleeping; first request takes 15–30s"],
              ["HMAC", "Keyed hash function for privacy-preserving identifier storage"],
          ])

  # Footer
    doc.add_page_break()
    heading(doc, "Document History", 1)
    table(doc, ["Version", "Date", "Changes"],
          [["1.0", date.today().strftime("%Y-%m-%d"),
            "Initial complete project documentation covering architecture, "
            "technology choices, hardware, software, 18 bugs, hardening, and proof results"]])

    try:
        doc.save(OUT)
        print(f"Saved: {OUT}")
    except PermissionError:
        alt = OUT.replace(".docx", "-v1.0.docx")
        doc.save(alt)
        print(f"Canonical locked. Saved: {alt}")


if __name__ == "__main__":
    build()
